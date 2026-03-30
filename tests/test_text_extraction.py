"""Tests for text extraction (extract_text in app.py)."""

import io

import pytest

# Import the cached function's underlying logic directly to avoid Streamlit context
from app import extract_text


def test_extract_txt():
    content = "Hello, this is a test policy document."
    file_bytes = content.encode("utf-8")
    result = extract_text(file_bytes, "policy.txt")
    assert result == content


def test_extract_txt_utf8():
    content = "Ärenden och åtgärder – GDPR-policy"
    file_bytes = content.encode("utf-8")
    result = extract_text(file_bytes, "gdpr.txt")
    assert "GDPR-policy" in result


def test_extract_unknown_extension_returns_empty():
    result = extract_text(b"some bytes", "file.xyz")
    assert result == ""


def test_extract_docx():
    """Create a minimal DOCX in-memory and verify extraction."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Second paragraph about GDPR compliance.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    result = extract_text(buf.read(), "test.docx")
    assert "test paragraph" in result
    assert "GDPR compliance" in result


def test_extract_pdf():
    """Create a minimal PDF using ReportLab and verify extraction."""
    try:
        import fitz  # noqa: F401
        from reportlab.pdfgen import canvas as rl_canvas
    except ImportError:
        pytest.skip("pymupdf or reportlab not installed")

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(100, 750, "Sample compliance policy document.")
    c.save()
    pdf_bytes = buf.getvalue()

    result = extract_text(pdf_bytes, "test.pdf")
    assert "compliance policy" in result
